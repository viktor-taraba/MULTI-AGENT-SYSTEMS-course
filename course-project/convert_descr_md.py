from docx import Document
from pathlib import Path
import os

def local_docx_to_markdown(input_filepath, output_filepath=None):
    try:
        doc = Document(input_filepath)
    except Exception as e:
        return f"Error loading document: {e}"
    
    md = []
    
    title = doc.paragraphs[0].text.strip() if doc.paragraphs else "Table Documentation"
    md.append(f"# {title}\n")
    
    for p in doc.paragraphs[1:]:
        text = p.text.strip()
        if text and text not in ["Columns", "Relations", "Unique keys", "Used by", "Triggers"]:
            md.append(f"> {text}\n")
            break
            
    for table in doc.tables:
        if not table.rows:
            continue
            
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        header_text = " ".join(header_cells).lower()
        
        # -- METADATA TABLE (Schema, Module, etc.) --
        if "documentation" in header_text or "schema" in header_text:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2 and cells[0]:
                    md.append(f"**{cells[0]}:** {cells[1]}")
            md.append("\n---")
            
        # -- STANDARD DATA TABLES (Columns, Relations, Keys) --
        elif any(keyword in header_text for keyword in ["data type", "primary table", "join", "key name"]):
            
             # Skip the Triggers table, which contains both 'key name' and 'when'
            if "when" in header_text:
                continue

            if "data type" in header_text:
                md.append("\n## Columns")
            elif "join" in header_text:
                md.append("\n## Relations")
            elif "key name" in header_text:
                md.append("\n## Unique Keys")

            md.append("| " + " | ".join(header_cells) + " |")
            md.append("|" + "|".join(["---"] * len(header_cells)) + "|")
            
            for row in table.rows[1:]:
                # Replace newlines inside cells so Markdown tables don't break
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                # Replace empty cells with a dash for cleaner Markdown
                cells = [c if c else "-" for c in cells]
                md.append("| " + " | ".join(cells) + " |")

    final_md = "\n".join(md)
    
    if output_filepath:
        with open(output_filepath, "w", encoding="utf-8") as f:
            f.write(final_md)
        print(f"Successfully saved to {output_filepath}")


def batch_convert_docs(input_folder: str, output_folder: str):
    """
    Iterates through a folder of .docx files and converts them to markdown.
    """

    input_path = Path(input_folder)
    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)
    
    if not input_path.exists():
        print(f"Error: The directory '{input_folder}' does not exist.")
        return

    docx_files = list(input_path.glob("*.docx"))
    
    if not docx_files:
        print(f"No .docx files found in '{input_folder}'.")
        return
        
    print(f"Found {len(docx_files)} files. Starting conversion...\n")

    for doc_file in docx_files:
        if doc_file.name.startswith("~$"):
            continue
        print(f"Converting: {doc_file.name}")
        output_filepath = output_path / f"{doc_file.stem}.md"
        local_docx_to_markdown(str(doc_file), str(output_filepath))
        
    print("\nAll files converted successfully!")

if __name__ == "__main__":
    SOURCE_FOLDER = "dwh_descr_docs"
    DESTINATION_FOLDER = "data"
    
    batch_convert_docs(SOURCE_FOLDER, DESTINATION_FOLDER)